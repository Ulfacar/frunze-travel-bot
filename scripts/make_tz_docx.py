"""Генерация ТЗ «Дырявое ведро / Ночной диспетчер» в Word (.docx) на Рабочий стол.

Собирает всё из сессии: диагноз, реальные прод-метрики (06.07), точки утечки,
решение «ночной диспетчер», фазовый план сборки, тех-риски, тесты, решения за
владельцем. Запуск: python scripts/make_tz_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = Path(r"C:\Users\alanb\OneDrive\Рабочий стол\Frunze-ТЗ-дырявое-ведро.docx")

NAVY = RGBColor(0x12, 0x2D, 0x42)
TEAL = RGBColor(0x00, 0x67, 0x6D)
AMBER = RGBColor(0x9A, 0x63, 0x00)
RED = RGBColor(0xB0, 0x3A, 0x2E)
SLATE = RGBColor(0x47, 0x56, 0x6A)
GREY = RGBColor(0x6B, 0x73, 0x85)


def _font(run, size, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def h1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run(text), 16, True, color)
    return p


def h2(doc, text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _font(p.add_run(text), 12.5, True, color)
    return p


def para(doc, text, size=11, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    _font(p.add_run(text), size, bold, color, italic)
    return p


def bullet(doc, text, size=11, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(text), size, False, color)
    return p


def numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(text), size)
    return p


def kv_table(doc, rows, headers=("Метрика", "Значение")):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"
    for j, hdr in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _font(c.paragraphs[0].add_run(hdr), 10.5, True, RGBColor(0xFF, 0xFF, 0xFF))
    for a, b in rows:
        r = t.add_row().cells
        r[0].text = ""; _font(r[0].paragraphs[0].add_run(a), 10.5, True, NAVY)
        r[1].text = ""; _font(r[1].paragraphs[0].add_run(b), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    for m in ("top", "bottom", "left", "right"):
        setattr(doc.sections[0], f"{m}_margin", Inches(0.8))

    # ---- титул ----
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ"), 11, True, AMBER)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("«Дырявое ведро»: ночной диспетчер и рост конверсии"), 22, True, NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    _font(p.add_run("Frunze Travel · GetVisa · Бишкек — 06.07.2026"), 11, False, SLATE)
    para(doc, "Тема: поток лидов есть, бот обрабатывает, но конверсия капает. Проблема не в "
              "рекламе (кран), а в дырах ведра. Задача — не увеличить поток, а перестать терять "
              "тех, за кого уже заплачено, и наконец видеть, где именно течёт.", 11, SLATE, italic=True)

    # ---- 1. диагноз ----
    h1(doc, "1. Диагноз")
    para(doc, "2 июля пришло больше 100 обращений, 4 менеджера не справляются. Воронка — «дырявое "
              "ведро»: бот консультирует → клиент доволен → уходит, в офис не идёт, цену унёс "
              "конкуренту. По визам обработка ровнее, по турам — шлют шаблон-простыню (консультация, "
              "а не продажа). Ключевой тезис владельца: «если с потоком не справляемся — нет смысла "
              "его увеличивать».")

    # ---- 2. реальные метрики ----
    h1(doc, "2. Реальные метрики (замер с прода 06.07)")
    para(doc, "Числа сняты по 389 боевым диалогам с запуска 01.07 (не прогнозы подрядчика).", 10.5, GREY)
    kv_table(doc, [
        ("Поток лидов/день", "медиана 66 (пик 96, минимум 22)"),
        ("Медиана ответа менеджера", "3 часа (182 мин) — ключевой факт утечки"),
        ("Лиды, рождённые ночью (22:00–08:00)", "20% (78 из 389)"),
        ("Застойные лиды (без человека, >2 суток)", "157 (визы 95, туры 62)"),
        ("Containment (бот ведёт сам)", "74% (289 из 389)"),
        ("Конверсия в продажу", "размечено 3 из 389 — НЕ измеряется"),
        ("Захват источника рекламы (CTWA)", "0% — атрибуция сломана"),
        ("Готовность (readiness-тиры)", "green 16% · warm 7% · noise 44%"),
    ])
    para(doc, "Вывод: главная утечка — время ответа (3 часа) и слепота (конверсию бизнес не считает: "
              "размечено 3 продажи из 389). «Вы не видите, где течёт» — буквально, не метафора.",
         11, RED, bold=True)

    # ---- 3. точки утечки ----
    h1(doc, "3. Точки утечки (9 точек, главная — ночь)")
    para(doc, "ГЛАВНАЯ — НОЧЬ: туры дозревают ночью (клиент дома, мечтает), а менеджеры в личном "
              "WhatsApp/спят. 20% лидов рождаются ночью, медиана ответа 3 часа → тёплый клиент "
              "остывает и уходит.", 11, AMBER, bold=True)
    for t in [
        "Три касания не делаются — нет системы «догнать».",
        "Консультация вместо продажи — бот дал цену, клиент ушёл к более дешёвому конкуренту.",
        "157 застойных лидов — оплаченный трафик лежит без касания.",
        "Билеты съедают 2–3 часа тур-менеджера на нерешительного клиента.",
        "Голосовые = сразу хендофф — войс висит, пока менеджер занят/спит.",
        "Нет слотов/календаря — некому перезвонить в назначенное время.",
        "Рекламный бюджет размазан поровну на неравные воронки.",
        "Две панели (своя админка + Bitrix) — двойная работа на каждой фиче.",
    ]:
        bullet(doc, t)

    # ---- 4. решение ----
    h1(doc, "4. Решение: «Ночной диспетчер»")
    para(doc, "Бот — НЕ продавец (клиентов на ботах не закрыть, цену-оружие демпинга не "
              "автоматизировать). Бот — ночной фильтр и диспетчер: не даёт лиду остыть и назначает "
              "человеческое касание.")
    h2(doc, "Что происходит ночью")
    for t in [
        "Клиент пишет в 23:47 — бот отвечает мгновенно, пока менеджер спит.",
        "Квалификация по 4 критериям: деньги · реальные даты · потребность · ЛПР (кто платит).",
        "Бронь слота утреннего звонка с окном ~30 минут.",
        "К 9:00 — досье менеджеру в Telegram: бюджет, состав семьи, направление, срочность; войс расшифрован.",
    ]:
        numbered(doc, t)
    h2(doc, "Разные KPI (снимают спор «бот убивает продажи»)")
    bullet(doc, "KPI бота — % слотов, где досье попало в цель (оружейник, не курьер).")
    bullet(doc, "KPI менеджера — % слотов, ставших деньгами.")
    bullet(doc, "Бонус: прозрачность стоимости лида по воронкам (первый честный счётчик).")

    # ---- 5. план внедрения ----
    h1(doc, "5. План внедрения (фазами, доказательство ценности до овербилда)")
    h2(doc, "Неделя 1 — «Досье в 9:00» (ноль состояния)")
    para(doc, "Cron + шаблон поверх готовой квалификации (readiness.py) + Telegram. Никакой слот-"
              "машины. Утром бот шлёт менеджеру карточку по каждому ночному лиду + ссылку на диалог. "
              "Самый быстрый «вау», ядро инфраструктуры на ~80% готово.")
    h2(doc, "Неделя 2 — «Утечка одним числом»")
    para(doc, "Одна колонка правды: по каждому досье — тронул менеджер лида до полудня или нет "
              "(считаем по факту переписки, менеджер ничего не нажимает). Раз в день владельцу: "
              "«Ночью 6 лидов, досье разослано 6, отработано 2, сгорело 4». Одна хардкод-эскалация "
              "(не тронут к N часам → пинг владельцу). Это доказывает утечку на живых данных.")
    h2(doc, "Ядро (только если утечка подтвердится) — TDD-инкременты")
    for t in [
        "Инкремент 1: сигнал ЛПР в readiness.py + функция night_qualified() (4 критерия), НЕ трогая tier(); регресс-тест-страховка.",
        "Инкремент 2: CallSlot как чистая стейт-машина (BOOKED→PUSHED→WORKED|BURNED→ESCALATED), переходы по timestamp+guard.",
        "Инкремент 3: персистентность (SlotStore Memory+Postgres, паритет одним contract-тестом).",
        "Инкремент 4: reconciliation on boot — просроченные слоты доигрываются при старте, без повторной отправки.",
        "Инкремент 5: идемпотентная утренняя джоба на scheduler.py (CAS mark_pushed → send; N тиков = 1 отправка).",
        "Инкременты 6–7: эскалация другому менеджеру + алерт владельцу; поимённое табло утечки.",
    ]:
        numbered(doc, t)

    # ---- 6. тех-требования и риски ----
    h1(doc, "6. Технические требования и риски")
    h2(doc, "Переиспользуем (готово ~60%)")
    for t in [
        "readiness.py — детерминированная квалификация (3 из 4 критериев; дописать ЛПР).",
        "scheduler.py — asyncio, тик 5 мин, реестр джоб (watchdog, автодожим).",
        "app/channels/telegram.py — канал для утреннего досье.",
        "Bitrix-зеркало (диалог→Lead + таймлайн) — отражать слот/исход комментарием.",
    ]:
        bullet(doc, t)
    h2(doc, "Риски и честные формулировки (в деке уже учтены)", color=RED)
    for t in [
        "Слот-стейт ОБЯЗАН переживать рестарт VPS (persist + reconciliation on boot) — иначе эскалация молча не выстрелит ночью. Главный тех-риск.",
        "Утренняя рассылка идемпотентна (guard/CAS) — без двойной отправки досье.",
        "«30 минут» — это окно на 5-мин тике, не таймер до секунды. Формулировать «в течение ~получаса».",
        "Голосовые ночью пока обрабатываем текстом; Whisper — отдельный шаг.",
        "«Честный счётчик $/лид» — прозрачность стоимости; выручка требует источника истины по сделкам (ручной исход / стадия Bitrix).",
        "MANAGERS-env fail-closed: пустой список = громкая ошибка на старте, не тихий [].",
    ]:
        bullet(doc, t, color=SLATE)

    # ---- 7. тестирование ----
    h1(doc, "7. Тестирование (risk-based)")
    para(doc, "Топ-риск: слот теряется на рестарте → эскалация молча не стреляет (тихий ночной "
              "провал). Первый тест — до единой строки логики диспетчера.")
    for t in [
        "test_reconcile_overdue_slot_escalates_once_after_reboot — главный (R1).",
        "Управляемое время (ManualClock), НИКАКИХ sleep — тик как функция времени.",
        "test_dossier_sent_exactly_once_on_repeated_tick + не пересылается после рестарта.",
        "test_escalation_boundary_invariant — escalated ⟺ now ≥ deadline, дрейф ≤ длины тика.",
        "test_empty_managers_fails_loud + эскалация без получателей → алерт владельцу.",
        "TourVisor юнитами НЕ крыть — только contract-тест адаптера на сохранённых ответах.",
        "Store-contract тест: Memory и Postgres через один набор переходов.",
    ]:
        bullet(doc, t)

    # ---- 8. решения за владельцем ----
    h1(doc, "8. Решения за владельцем (гейтят код)")
    for t in [
        "Цена туров: бот даёт вилку/тизер, точную — человек? (защищает схему демпинга).",
        "Primary-панель: живём в Bitrix или в своей админке? (убрать двойной налог).",
        "Билеты: выделяем одного менеджера (двое — на туры)? Рекламу оставляем или снимаем?",
        "Календарь: встроенный + Telegram-пуши или Google Calendar?",
        "Ре-энгейдж: бот пишет клиентам сам (риск бана WhatsApp) или только задачи менеджерам? (рекомендация — задачи).",
        "Воронки в Bitrix уже разделены туры/визы/билеты или всё в общей?",
        "Рекламный бюджет — уточнить сумму и период (день/месяц).",
    ]:
        numbered(doc, t)

    # ---- 9. вне скоупа сейчас ----
    h1(doc, "9. Отдельные задачи (вне скоупа ночного диспетчера)")
    for t in [
        "Прод-пожары: пополнить OpenRouter, держать оплаченным TourVisor (лечится деньгами, не кодом).",
        "Атрибуция CTWA: захват 0% — залогировать сырые payload'ы Wappi, поймать формат, поправить парсер. Дешёвый диагностический заход.",
        "Билеты: перевод на self-service ботом ИЛИ выделенный менеджер (бизнес-решение).",
        "Перевод менеджеров с личного WhatsApp в Bitrix + обучалка по продажам (Кирилл).",
    ]:
        bullet(doc, t)

    para(doc, "")
    para(doc, "Материалы к ТЗ: презентация Frunze-дырявое-ведро.pptx (Рабочий стол) · "
              "интерактивная HTML-дека · сценарий живого показа docs/pitch-script-leaky-bucket.md.",
         9.5, GREY, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK → {OUT}")


if __name__ == "__main__":
    build()
