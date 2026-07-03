from __future__ import annotations

import argparse
import re
import shutil
import time
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SECTION_CONFIGS = {
    "visas": {
        "header_marker": "ВИЗЫ",
        "label": "ВИЗАМ",
        "output_name": "Frunze-FAQ-ВИЗЫ.docx",
    },
    "tours": {
        "header_marker": "ТУРЫ",
        "label": "ТУРАМ",
        "output_name": "Frunze-FAQ-ТУРЫ.docx",
    },
    "tickets": {
        "header_marker": "АВИАБИЛЕТЫ",
        "label": "АВИАБИЛЕТАМ",
        "output_name": "Frunze-FAQ-БИЛЕТЫ.docx",
    },
}

INSTRUCTION = (
    "Под каждым вопросом впишите короткий ответ, как отвечаете клиенту. "
    "Эти ответы бот будет давать мгновенно и бесплатно. "
    "Без гарантий одобрения визы, без обещаний."
)

FOOTER_TEXT = "Заполненный файл верните — вопросы станут ответами бота."


def default_desktop() -> Path:
    return Path(__file__).resolve().parents[2]


def default_source_path() -> Path:
    return default_desktop() / "Frunze-вопросы-клиентов.md"


def parse_questions(markdown_text: str) -> dict[str, list[tuple[int, str]]]:
    headers = list(re.finditer(r"^#\s+(.+?)\s*$", markdown_text, flags=re.MULTILINE))
    sections: dict[str, list[tuple[int, str]]] = {}

    for index, header in enumerate(headers):
        title = header.group(1)
        start = header.end()
        end = headers[index + 1].start() if index + 1 < len(headers) else len(markdown_text)
        body = markdown_text[start:end]

        for key, config in SECTION_CONFIGS.items():
            if config["header_marker"] in title:
                questions = []
                for match in re.finditer(r"^\s*(\d+)\.\s+(.+?)\s*$", body, flags=re.MULTILINE):
                    number = int(match.group(1))
                    text = match.group(2).strip()
                    if text:
                        questions.append((number, text))
                sections[key] = questions

    missing = [key for key in SECTION_CONFIGS if key not in sections]
    if missing:
        missing_labels = ", ".join(SECTION_CONFIGS[key]["header_marker"] for key in missing)
        raise ValueError(f"Не найдены секции: {missing_labels}")

    return sections


def set_run_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def copy_docx(temp_path: Path, path: Path) -> None:
    if not path.exists():
        shutil.copyfile(temp_path, path)
        return

    backup_path = Path(__file__).resolve().parents[1] / f"backup-frunze-faq-{uuid.uuid4().hex}.docx"
    shutil.copyfile(path, backup_path)
    try:
        for attempt in range(5):
            try:
                path.unlink()
                break
            except PermissionError:
                if attempt == 4:
                    raise
                time.sleep(0.4)

        for attempt in range(5):
            try:
                shutil.copyfile(temp_path, path)
                break
            except PermissionError:
                if attempt == 4:
                    raise
                time.sleep(0.4)
    except Exception:
        if not path.exists() and backup_path.exists():
            shutil.copyfile(backup_path, path)
        if path.exists() and len(Document(path).paragraphs) > 0:
            return
        raise
    finally:
        if backup_path.exists():
            backup_path.unlink()


def build_docx(path: Path, label: str, questions: list[tuple[int, str]]) -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title_run = title.add_run(f"Frunze Travel — вопросы клиентов по {label}. Впишите ответы для бота")
    set_run_font(title_run, size=14, bold=True)
    title.paragraph_format.space_after = Pt(10)

    instruction = document.add_paragraph()
    instruction_run = instruction.add_run(INSTRUCTION)
    set_run_font(instruction_run, size=11)
    instruction.paragraph_format.space_after = Pt(12)

    for number, question in questions:
        question_paragraph = document.add_paragraph()
        question_run = question_paragraph.add_run(f"{number}. {question}")
        set_run_font(question_run, size=11, bold=True)
        question_paragraph.paragraph_format.space_after = Pt(3)

        answer_paragraph = document.add_paragraph()
        answer_run = answer_paragraph.add_run("Ответ: ____________________________")
        set_run_font(answer_run, size=11)
        answer_paragraph.paragraph_format.space_after = Pt(8)

    closing = document.add_paragraph()
    closing_run = closing.add_run(FOOTER_TEXT)
    set_run_font(closing_run, size=11)

    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = Path(__file__).resolve().parents[1] / f"tmp-frunze-faq-{uuid.uuid4().hex}.docx"
    try:
        document.save(temp_path)
        copy_docx(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def make_docs(source_path: Path, output_dir: Path) -> dict[str, tuple[Path, int]]:
    markdown_text = source_path.read_text(encoding="utf-8")
    sections = parse_questions(markdown_text)

    results: dict[str, tuple[Path, int]] = {}
    for key, questions in sections.items():
        config = SECTION_CONFIGS[key]
        output_path = output_dir / config["output_name"]
        for existing_path in output_dir.glob("Frunze-FAQ-*.docx"):
            if existing_path.name.casefold() == config["output_name"].casefold():
                output_path = existing_path
                break
        build_docx(output_path, config["label"], questions)
        results[key] = (output_path, len(questions))

    return results


def main() -> None:
    parser = argparse.ArgumentParser(description="Create Frunze Travel FAQ questionnaires as DOCX files.")
    parser.add_argument("--source", type=Path, default=default_source_path())
    parser.add_argument("--output-dir", type=Path, default=default_desktop())
    args = parser.parse_args()

    results = make_docs(args.source, args.output_dir)
    for key in SECTION_CONFIGS:
        path, count = results[key]
        print(f"{path}\t{count}")


if __name__ == "__main__":
    main()
